from __future__ import annotations

import hashlib
from pathlib import Path

from app.models import RawDocument
from app.utils.text import compact_whitespace


def _import_document():
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency check
        raise RuntimeError(
            "python-docx is not installed. Install dependencies from requirements.txt."
        ) from exc
    return Document


def read_docx(path: Path) -> RawDocument:
    document_factory = _import_document()
    document = document_factory(str(path))
    paragraphs = []

    for paragraph in document.paragraphs:
        text = compact_whitespace(paragraph.text)
        if text:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [compact_whitespace(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)
    sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
    return RawDocument(
        source_path=str(path),
        file_name=path.name,
        paragraphs=paragraphs,
        full_text=full_text,
        sha256=sha256,
    )
